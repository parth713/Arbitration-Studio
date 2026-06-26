from dataclasses import dataclass
from io import BytesIO
import re
from typing import Dict, Iterable, List, Tuple

from docx import Document as DocxDocument
from pypdf import PdfReader


@dataclass
class PageText:
    page_number: int
    text: str


@dataclass
class SourceDocument:
    doc_id: str
    filename: str
    kind: str
    pages: List[PageText]
    confidence: int = 0
    evidence: str = ""

    @property
    def full_text(self) -> str:
        return "\n\n".join(page.text for page in self.pages)


def extract_file(uploaded_file) -> List[PageText]:
    name = uploaded_file.name.lower()
    data = uploaded_file.getvalue()
    if name.endswith(".pdf"):
        return _extract_pdf(data)
    if name.endswith(".docx"):
        return _extract_docx(data)
    raise ValueError(f"Unsupported file type: {uploaded_file.name}")


def classify_document(filename: str, text: str) -> str:
    return classify_document_detailed(filename, text)[0]


def classify_document_detailed(filename: str, text: str) -> Tuple[str, int, str]:
    normalized_filename = _normalize(filename)
    full_text = _normalize(text[:20000])
    opening = _normalize(text[:3500])
    first_lines = _normalize("\n".join(line.strip() for line in text.splitlines()[:30] if line.strip()))

    scores: Dict[str, int] = {
        "Statement of Claim": 0,
        "Statement of Defence": 0,
        "Rejoinder": 0,
    }
    evidence: Dict[str, List[str]] = {kind: [] for kind in scores}

    _apply_patterns(scores, evidence, "Statement of Claim", normalized_filename, [
        (r"\bsoc\b", 5, "filename contains SOC"),
        (r"\bstatement\s+of\s+claim\b", 5, "filename says statement of claim"),
        (r"\bpoints\s+of\s+claim\b", 5, "filename says points of claim"),
    ])
    _apply_patterns(scores, evidence, "Statement of Defence", normalized_filename, [
        (r"\bsod\b", 6, "filename contains SOD"),
        (r"\bstatement\s+of\s+defen[cs]e\b", 7, "filename says statement of defence"),
        (r"\bdefen[cs]e\b", 3, "filename says defence"),
        (r"\bwritten\s+statement\b", 3, "filename says written statement"),
        (r"\bcounterclaim\b", 3, "filename says counterclaim"),
    ])
    _apply_patterns(scores, evidence, "Rejoinder", normalized_filename, [
        (r"\brejoinder\b", 7, "filename says rejoinder"),
        (r"\breply\b", 3, "filename says reply"),
        (r"\breplication\b", 5, "filename says replication"),
    ])

    _apply_patterns(scores, evidence, "Statement of Claim", first_lines, [
        (r"\bstatement\s+of\s+claim\b", 10, "opening title says statement of claim"),
        (r"\bpoints\s+of\s+claim\b", 10, "opening title says points of claim"),
        (r"\bclaimant'?s?\s+statement\s+of\s+claim\b", 10, "opening claimant SOC title"),
    ])
    _apply_patterns(scores, evidence, "Statement of Defence", first_lines, [
        (r"\bstatement\s+of\s+defen[cs]e\b", 12, "opening title says statement of defence"),
        (r"\bstatement\s+of\s+defen[cs]e\s+(?:and|&)\s+counterclaim\b", 14, "opening title says defence and counterclaim"),
        (r"\bdefen[cs]e\s+to\s+(?:the\s+)?statement\s+of\s+claim\b", 12, "opening says defence to SOC"),
        (r"\brespondent'?s?\s+statement\s+of\s+defen[cs]e\b", 12, "opening respondent SOD title"),
    ])
    _apply_patterns(scores, evidence, "Rejoinder", first_lines, [
        (r"\brejoinder\b", 12, "opening title says rejoinder"),
        (r"\breply\s+to\s+(?:the\s+)?statement\s+of\s+defen[cs]e\b", 12, "opening says reply to SOD"),
        (r"\bclaimant'?s?\s+reply\b", 8, "opening claimant reply title"),
        (r"\breplication\b", 10, "opening title says replication"),
    ])

    _apply_patterns(scores, evidence, "Statement of Claim", opening, [
        (r"\bclaimant\s+(?:submits|states|avers|claims)\b", 3, "claimant asserts claims"),
        (r"\bcause\s+of\s+action\b", 3, "cause of action language"),
        (r"\breliefs?\s+sought\b", 4, "reliefs sought section"),
        (r"\bprayer\s+for\s+relief\b", 3, "prayer for relief"),
    ])
    _apply_patterns(scores, evidence, "Statement of Defence", opening, [
        (r"\brespondent\s+(?:denies|submits|states|avers)\b", 5, "respondent denies/submits"),
        (r"\bden(?:y|ies|ied)\s+(?:each|all|the)\s+(?:allegation|claim|averment)", 5, "denial language"),
        (r"\bwithout\s+prejudice\b", 2, "without prejudice defence language"),
        (r"\bpreliminary\s+objections?\b", 4, "preliminary objections"),
        (r"\bcounterclaim\b", 5, "counterclaim language"),
        (r"\bparagraph[-\s]?wise\s+reply\b", 6, "paragraph-wise reply"),
    ])
    _apply_patterns(scores, evidence, "Rejoinder", opening, [
        (r"\bclaimant\s+(?:denies|replies|responds)\b", 4, "claimant replies/denies"),
        (r"\bin\s+rejoinder\b", 5, "in rejoinder language"),
        (r"\breply\s+to\s+(?:respondent'?s?\s+)?(?:statement\s+of\s+)?defen[cs]e\b", 7, "reply to defence language"),
    ])

    _apply_patterns(scores, evidence, "Statement of Defence", full_text, [
        (r"\bdefen[cs]e\s+to\s+(?:the\s+)?statement\s+of\s+claim\b", 10, "document responds to SOC"),
        (r"\bthe\s+statement\s+of\s+claim\s+is\s+(?:denied|misconceived|untenable)", 7, "SOC is denied"),
        (r"\ballegations?\s+in\s+(?:the\s+)?statement\s+of\s+claim\s+(?:are|is)\s+denied\b", 7, "allegations in SOC denied"),
    ])
    _apply_patterns(scores, evidence, "Rejoinder", full_text, [
        (r"\bthe\s+statement\s+of\s+defen[cs]e\s+is\s+(?:denied|misconceived|untenable)", 8, "SOD is denied"),
        (r"\brespondent'?s?\s+defen[cs]e\s+is\s+denied\b", 7, "respondent defence denied"),
    ])

    # A defence naturally mentions the Statement of Claim many times. Treat isolated SOC
    # references as background unless they appear as an opening title or filename signal.
    if scores["Statement of Defence"] >= 10 and scores["Statement of Claim"] <= scores["Statement of Defence"] + 4:
        scores["Statement of Claim"] = max(0, scores["Statement of Claim"] - 6)
        evidence["Statement of Defence"].append("SOC references treated as response context")

    if scores["Rejoinder"] >= 10 and scores["Statement of Defence"] <= scores["Rejoinder"] + 4:
        scores["Statement of Defence"] = max(0, scores["Statement of Defence"] - 5)
        evidence["Rejoinder"].append("SOD references treated as response context")

    best, best_score = max(scores.items(), key=lambda item: item[1])
    ordered = sorted(scores.items(), key=lambda item: item[1], reverse=True)
    runner_up = ordered[1][1]
    if best_score < 5 or best_score - runner_up < 2:
        return "Supporting Document", best_score, _evidence_text(scores, evidence, "Supporting Document")
    return best, best_score, _evidence_text(scores, evidence, best)


def make_source_documents(uploaded_files: Iterable) -> List[SourceDocument]:
    documents = []
    for idx, uploaded_file in enumerate(uploaded_files, start=1):
        pages = extract_file(uploaded_file)
        text = "\n".join(page.text for page in pages)
        kind, confidence, evidence = classify_document_detailed(uploaded_file.name, text)
        documents.append(
            SourceDocument(
                doc_id=f"D{idx}",
                filename=uploaded_file.name,
                kind=kind,
                pages=pages,
                confidence=confidence,
                evidence=evidence,
            )
        )
    return documents


def _extract_pdf(data: bytes) -> List[PageText]:
    reader = PdfReader(BytesIO(data))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        pages.append(PageText(page_number=index, text=(page.extract_text() or "").strip()))
    return pages


def _extract_docx(data: bytes) -> List[PageText]:
    doc = DocxDocument(BytesIO(data))
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
    return [PageText(page_number=1, text=text.strip())]


def _normalize(text: str) -> str:
    text = text.lower()
    text = re.sub(r"[_\-]+", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _apply_patterns(
    scores: Dict[str, int],
    evidence: Dict[str, List[str]],
    kind: str,
    text: str,
    patterns: List[Tuple[str, int, str]],
) -> None:
    for pattern, weight, reason in patterns:
        matches = re.findall(pattern, text, flags=re.IGNORECASE)
        if matches:
            capped_matches = min(len(matches), 3)
            scores[kind] += weight * capped_matches
            evidence[kind].append(reason)


def _evidence_text(scores: Dict[str, int], evidence: Dict[str, List[str]], selected: str) -> str:
    score_text = ", ".join(f"{kind}: {score}" for kind, score in sorted(scores.items()))
    if selected == "Supporting Document":
        return f"Insufficient pleading-specific signals. Scores: {score_text}"
    reasons = "; ".join(dict.fromkeys(evidence[selected])) or "No evidence captured"
    return f"{reasons}. Scores: {score_text}"
